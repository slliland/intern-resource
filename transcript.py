from docx import Document
"""
拆分访谈笔录
"""
# 打开.docx文件
document = Document('/Users/songyujian/Downloads/Material Used for AI-BotV/Concept material/Whitening Essence_Age 20-30/Transcript-G4 Non Essence User.docx')

# 创建一个新的Document对象来存储过滤后的内容
filtered_document = Document()

# 初始化一个标志变量，用于跟踪是否需要跳过下一段落
skip_next_para = False

for i, para in enumerate(document.paragraphs):
    text = para.text.strip()

    # 如果当前段落应该被跳过，则重置标志并继续循环
    if skip_next_para:
        skip_next_para = False
        continue

    # 检查当前段落是否以"M："开头且下一段落以"1/2："开头（写哪个就跳过那个人）
    if text.startswith("M：") and (
            i + 1 < len(document.paragraphs) and document.paragraphs[i + 1].text.strip().startswith("1：")):
        # 设置标志以跳过下一段落（即"2："）
        skip_next_para = True
        continue

    # 当前段落不是针对"2："的提问也不是"2："回答，将其添加到新文档中
    filtered_document.add_paragraph(para.text)

# 保存过滤后的文档
filtered_document.save('/Users/songyujian/Downloads/Essence拆分结果/邓KY Practical Seeker, Non-Essence User.docx')
