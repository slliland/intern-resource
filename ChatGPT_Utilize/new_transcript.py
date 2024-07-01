# -*- coding: utf-8 -*-
"""
用于AI总结访谈笔录
"""
import os
from docx import Document
import pandas as pd
import openai
import re
from concurrent.futures import ThreadPoolExecutor, as_completed

# 设置 OpenAI API 密钥和基础 URL
openai.api_key = os.getenv('OPENAI_API_KEY')
openai.api_base = os.getenv("OPENAI_API_BASE")

# 与 OpenAI 聊天完成 API 交互的函数
def openai_create_chat_completions(prompt, content):
    try:
        completion = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": content}
            ],
            app_id="uOv9nxRzGmJtONU5",  # 必填,申请后提供
            app_secret="Rc9jy3Tdo0teTUrVYSv6rRf7EvMTvoeW",  # 必填,申请后提供
            user_email="songyj@foxmail.com",  # 必填,自己的邮箱
            tag="2400159201"  # 选填，应用或项目名称
        )
        return completion.choices[0].message['content']
    except openai.error.InvalidRequestError as e:
        print(f"请求被过滤：{e}")
        return None

"""
非并行版本
"""
# def split_and_process_docx(file_path, output_dir, pages_per_file=20):
#     # 加载 Word 文档
#     doc = Document(file_path)
#     # 初始化变量
#     part_num = 1
#     text_blocks = []
#     data_for_excel = []
#     for i, paragraph in enumerate(doc.paragraphs):
#         text_blocks.append(paragraph.text)
#
#         # 检查是否应该拆分文件（每二十段或文档末尾）
#         if (i + 1) % pages_per_file == 0 or i == len(doc.paragraphs) - 1:
#             # 将当前部分内容合并为字符串，并交给 AI 处理
#             content_to_process = "\n".join(text_blocks)
#             # 构建 prompt（提示）
#             prompt = "以下是一个抗老护肤品入户深访笔录的部分信息，请基于M的提问内容和这个被访者的回答内容，帮我做一个内容汇总，把被访者回答出的所有信息根据不同主题归纳成一整段话，用第一人称，需要非常全面且具体，任何被访者说的话都不要遗漏。" \
#                      "最终给到我格式严格规定为：1、根据归纳的主题产生的问句（问句前面的序号必须为1、）  2、被访者回答出的对应内容（使用原话，回答前面的序号必须为2、）。严格规定不得使用其他的回复格式" \
#                      "回答示例：" \
#                      "1、你们俩都介绍一下自己，让我们熟悉一点，聊起来更通畅一点，好不好？" \
#                      "2、我叫小胡，今年是21岁，是在读大学生，大三。平时喜欢在课外去看一些演唱会，一些现场的音乐那些。"
#             ai_response = openai_create_chat_completions(prompt, content_to_process)
#
#             if ai_response is None:
#                 print(f"由于内容管理策略，跳过当前段落处理：{content_to_process}")
#                 continue
#
#             # 打印 AI 响应结果以便调试
#             print(ai_response)
#
#             # 使用正则表达式匹配多个 "1、问题 2、答案" 结构
#             pattern = r'1[、. ：:](.*?)\s*2[、. ：:](.*?)(?=\s*1[、. ：:]|\Z)'
#             matches = re.finditer(pattern, ai_response, re.DOTALL)
#
#             for match in matches:
#                 question, answer = match.groups()
#                 data_for_excel.append({'Question': question.strip(), 'Answer': answer.strip()})
#
#             if not matches:
#                 print(f"无法解析或未检测到完整问题编号：{ai_response}")
#
#             # 准备处理下一个文档部分
#             text_blocks.clear()
#             print("当前为第" + str(part_num) + "轮")
#             part_num += 1
#
#     df = pd.DataFrame(data_for_excel)
#     excel_filename = f"{output_dir}/interview_summary.xlsx"
#     df.to_excel(excel_filename, index=False)

"""
并行版本
"""
def process_block(content_to_process, prompt):
    # 构建 prompt 并调用 OpenAI 的 API
    ai_response = openai_create_chat_completions(prompt, content_to_process)
    return ai_response


def split_and_process_docx(file_path, output_dir, pages_per_file=4):
    # 加载 Word 文档
    doc = Document(file_path)
    # 初始化变量
    data_for_excel = []

    # 使用 ThreadPoolExecutor 创建线程池
    with ThreadPoolExecutor() as executor:
        futures = []

        for i in range(0, len(doc.paragraphs), pages_per_file):
            text_blocks = doc.paragraphs[i:i + pages_per_file]
            content_to_process = "\n".join(paragraph.text for paragraph in text_blocks)
            prompt = "以下是一个抗老护肤品入户深访笔录的部分信息，请基于M的提问内容和这个被访者的回答内容，帮我做一个内容汇总，把被访者回答出的所有信息根据不同主题归纳成一整段话，用第一人称，需要非常全面且具体，任何被访者说的话都不要遗漏。" \
                                     "最终给到我格式严格规定为：1、根据归纳的主题产生的问句（问句前面的序号严格规定必须为1、）  2、被访者回答出的对应内容（使用原话，回答前面的序号严格规定必须为2、）。严格规定不得使用其他的回复格式" \
                                 "回答示例：" \
                                 "1、你们俩都介绍一下自己，让我们熟悉一点，聊起来更通畅一点，好不好？" \
                                 "2、我叫小胡，今年是21岁，是在读大学生，大三。平时喜欢在课外去看一些演唱会，一些现场的音乐那些。"

            # 提交任务到线程池
            futures.append(executor.submit(process_block, content_to_process, prompt))

        for future in as_completed(futures):
            ai_response = future.result()

            if ai_response is None:
                print(f"由于内容管理策略，跳过当前段落处理")
                continue

            # 打印 AI 响应结果以便调试
            print(ai_response)

            # 使用正则表达式匹配多个 "1、问题 2、答案" 结构
            pattern = r'1[、. ：:](.*?)\s*2[、. ：:](.*?)(?=\s*1[、. ：:]|\Z)'
            matches = re.finditer(pattern, ai_response, re.DOTALL)

            for match in matches:
                question, answer = match.groups()
                data_for_excel.append({'Question': question.strip(), 'Answer': answer.strip()})

            if not matches:
                print(f"无法解析或未检测到完整问题编号：{ai_response}")

    df = pd.DataFrame(data_for_excel)
    excel_filename = f"{output_dir}/{obj} interview_summary.xlsx"
    df.to_excel(excel_filename, index=False)


# 设置原始文件路径和输出目录
obj = "邓KY Practical Seeker, Non-Essence User"
file_path = "/Users/songyujian/Downloads/Essence拆分结果/"+obj+".docx"
# file_path = "/Users/songyujian/Downloads/测试文件.docx"
output_dir = "/Users/songyujian/Downloads/AI提取"

# 确保输出目录存在，如果不存在则创建它
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# 调用函数执行拆分和处理操作
split_and_process_docx(file_path, output_dir)
