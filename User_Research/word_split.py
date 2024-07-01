import os
from docx import Document


def split_docx(file_path, output_dir, number_of_parts=8):
    # 确保文件存在
    if not os.path.exists(file_path):
        print(f"File does not exist: {file_path}")
        return

    # 加载文档
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Could not open file {file_path}: {e}")
        return

    # 获取所有非空段落
    paragraphs = [p for p in doc.paragraphs if p.text.strip() != '']

    # 计算每份文档应有的段落数量
    part_length = len(paragraphs) // number_of_parts + (len(paragraphs) % number_of_parts > 0)

    # 拆分文档并保存
    for i in range(number_of_parts):
        # 创建新文档
        new_doc = Document()

        # 计算当前部分的起始和结束索引
        start_index = i * part_length
        end_index = start_index + part_length

        # 将当前部分的段落添加到新文档中
        for paragraph in paragraphs[start_index:end_index]:
            new_doc.add_paragraph(paragraph.text)

        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

        # 保存新文档到指定的输出目录下
        new_file_path = os.path.join(output_dir, f"part_{i + 1}_{os.path.basename(file_path)}")
        new_doc.save(new_file_path)


# 指定输入和输出目录路径
input_directory_path = '/Users/songyujian/Downloads/深访_精华眼霜用户研究/input'
output_base_directory_path = '/Users/songyujian/Downloads/深访_精华眼霜用户研究/output'

# 遍历目录下所有Word文档 (.docx)
for filename in os.listdir(input_directory_path):
    if filename.endswith('.docx'):
        file_path = os.path.join(input_directory_path, filename)

        # 创建以文件名命名的输出目录
        output_dir = os.path.join(output_base_directory_path, os.path.splitext(filename)[0])

        # 确保文件路径正确且为文件而不是目录
        if os.path.isfile(file_path):
            split_docx(file_path, output_dir)
            print(f"完成拆分：{filename}")
        else:
            print(f"Skipped non-existing or invalid file: {file_path}")

print("所有文档拆分完成！")
